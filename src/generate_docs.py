from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime
from pathlib import Path

def create_product_research(json_path, screenshots_dir, output_file):
    """Generate a Word document with the product research"""
    # Create a new Document
    doc = Document()
    
    # Read the JSON data
    with open(json_path, 'r') as f:
        json_data = json.load(f)
    
    # Title
    title = doc.add_heading(f"{json_data['product_name']} Research Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add category
    doc.add_paragraph(f"Category: {json_data['product_category']}")
    
    # Add date
    doc.add_paragraph(f"Research Date: {datetime.now().strftime('%Y-%m-%d')}")
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    # Remove video references from executive summary
    summary = json_data['executive_summary'].replace("video demonstrates", "document covers")
    summary = summary.replace("video", "analysis")
    doc.add_paragraph(summary)
    
    # Product Overview
    doc.add_heading('Product Overview', level=1)
    overview = json_data['product_overview']
    doc.add_paragraph(overview['description'])
    
    # Target Audience
    doc.add_heading('Target Audience', level=2)
    doc.add_paragraph(overview['target_audience'])
    
    # Key Features
    doc.add_heading('Key Features', level=2)
    for feature in overview['key_features']:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Workflow Analysis
    doc.add_heading('Workflow Analysis', level=1)
    
    # Sort screenshots by timestamp
    screenshots_dir = Path(screenshots_dir)
    screenshots = list(screenshots_dir.glob('*.png'))
    screenshots.sort()
    
    for step in json_data['workflow_steps']:
        # Skip non-major steps to focus on key interactions
        if not step['is_major_step']:
            continue
            
        # Create step heading without timestamp
        prefix = "Major Step: " if step['is_major_step'] else "Step: "
        step_desc = step['description']
        # Remove any timestamp references from description
        step_desc = ' '.join([s for s in step_desc.split() if ':' not in s and 'timestamp' not in s.lower()])
        doc.add_heading(step_desc, level=2)
        
        # Add UI Elements
        if step['ui_elements']:
            doc.add_heading('UI Elements', level=3)
            for element in step['ui_elements']:
                doc.add_paragraph(element, style='List Bullet')
        
        # Add User Interaction
        doc.add_heading('User Interaction', level=3)
        interaction = step['user_interaction']
        # Remove any timestamp or video references
        interaction = interaction.replace("video", "document")
        interaction = ' '.join([s for s in interaction.split() if ':' not in s])
        doc.add_paragraph(interaction)
        
        # Add Design Analysis
        doc.add_heading('Design Analysis', level=3)
        doc.add_paragraph(step['design_analysis'])
        
        # Add Technical Observations
        doc.add_heading('Technical Observations', level=3)
        doc.add_paragraph(step['technical_observations'])
        
        # Add screenshot if available
        timestamp = step['timestamp'].replace(':', '_')
        matching_screenshots = [s for s in screenshots if timestamp in s.name]
        if matching_screenshots:
            doc.add_picture(str(matching_screenshots[0]), width=Inches(6))
            doc.add_paragraph()  # Add spacing after image
    
    # Key Findings
    doc.add_heading('Key Findings', level=1)
    findings = json_data['key_findings']
    
    # Usability Insights
    doc.add_heading('Usability Insights', level=2)
    for insight in findings['usability_insights']:
        doc.add_paragraph(insight, style='List Bullet')
    
    # Design Patterns
    doc.add_heading('Design Patterns', level=2)
    for pattern in findings['design_patterns']:
        doc.add_paragraph(pattern, style='List Bullet')
    
    # Technical Highlights
    doc.add_heading('Technical Highlights', level=2)
    for highlight in findings['technical_highlights']:
        doc.add_paragraph(highlight, style='List Bullet')
    
    # Save the document
    output_file = Path(output_file)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    print(f"Product research document saved to {output_file}")

if __name__ == "__main__":
    # Load the JSON data
    json_path = 'video_analysis_output.json'
    screenshots_dir = "screenshot_output"
    output_file = "product_research.docx"
    
    create_product_research(json_path, screenshots_dir, output_file)
